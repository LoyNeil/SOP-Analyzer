import ssl
import asyncio
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv
from openai import AsyncOpenAI
import httpx
import os
import json

load_dotenv()

ssl_cert = os.getenv("SSL_CERT_FILE")
ssl_context = ssl.create_default_context()
ssl_context.load_verify_locations(cafile=ssl_cert)
ssl_context.verify_flags &= ~ssl.VERIFY_X509_STRICT
http_client = httpx.AsyncClient(verify=ssl_context)

client = AsyncOpenAI(
    api_key=os.getenv("OPENAI_API_KEY"),
    http_client=http_client
)

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ─────────────────────────────────────────
#  FILE PARSERS
# ─────────────────────────────────────────

async def extract_text_from_file(file: UploadFile) -> str:
    """
    Single extraction function used by BOTH /analyse and /analyse-sop.
    Handles .vtt, .txt  → raw UTF-8 decode  (Zoom transcripts)
             .docx      → python-docx paragraph extraction
             .pdf       → pdfplumber text extraction
    Calls file.seek(0) first so it's safe even if the stream was partially read.
    """
    await file.seek(0)
    content = await file.read()
    filename = (file.filename or "").lower()

    if filename.endswith(".docx"):
        return _extract_docx(content)

    if filename.endswith(".pdf"):
        return _extract_pdf(content)

    # Default: plain text / vtt
    return content.decode("utf-8", errors="ignore")


def _extract_docx(content: bytes) -> str:
    """Extract all paragraph text from a .docx file, preserving order."""
    try:
        from docx import Document
        import io
        doc = Document(io.BytesIO(content))
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        # Also pull text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "  |  ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    lines.append(row_text)
        return "\n".join(lines)
    except ImportError:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        )
    except Exception as e:
        raise RuntimeError(f"Failed to parse .docx file: {e}")


def _extract_pdf(content: bytes) -> str:
    """Extract text from a text-based PDF."""
    try:
        import pdfplumber
        import io
        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)
    except ImportError:
        raise RuntimeError(
            "pdfplumber is not installed. Run: pip install pdfplumber"
        )
    except Exception as e:
        raise RuntimeError(f"Failed to parse .pdf file: {e}")


# ─────────────────────────────────────────
#  AI CALLER
# ─────────────────────────────────────────

async def call_openai(prompt: str, model: str = "gpt-5.2") -> str:
    response = await client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are an expert business analyst who specialises in "
                    "Standard Operating Procedures (SOPs). "
                    "Always respond in clean, structured JSON only — no markdown, no extra text."
                )
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        temperature=0,
    )
    return response.choices[0].message.content


def clean_json(result: str) -> dict:
    result = result.strip()
    if result.startswith("```"):
        result = result.split("\n", 1)[1]
        result = result.rsplit("```", 1)[0]
    return json.loads(result)


# ─────────────────────────────────────────
#  TRANSCRIPT PROMPTS  (unchanged)
# ─────────────────────────────────────────

def build_sop_prompt(transcript: str) -> str:
    return f"""
You are a process documentation expert. Analyze the transcript and return ONLY this JSON — no extra text, no markdown fences.

{{
  "sop": {{
    "title": "SOP title derived from the transcript topic",
    "generated_from": "short description of the meeting",
    "date": "today's date in MM/DD/YYYY format",
    "version": "1.0",
    "participants": 0,
    "tags": ["tag1", "tag2"],
    "purpose": "2–3 sentence paragraph explaining the purpose of this SOP.",
    "scope": "2–3 sentence paragraph describing what this SOP applies to.",
    "roles": [
      {{
        "name": "Role name",
        "responsibility": "Full description of everything this role does across the process."
      }}
    ],
    "prerequisites": [
      "Specific access, knowledge, or condition required before the process begins."
    ],
    "start_state": "The exact condition that triggers this process to begin.",
    "end_state": "The exact condition that confirms the process is complete.",
    "steps": [
      {{
        "number": 1,
        "title": "Step title — concise verb phrase",
        "action": "Name the role at the start. Describe exactly what is done and which tool is used.",
        "outcome_or_decision": "For DECISION steps: phrase as a Yes/No question. For PROCESS steps: describe the output produced.",
        "happy_path": "Successful outcome. End with: Continue to **Step N – Title**. Final step ends with: Process completed successfully.",
        "unhappy_path": "Failure outcome. End with: Return to **Step N – Title**. or: remains in **Step N – Title**. Write N/A only if genuinely no failure path exists.",
        "handoffs": "Format as [Role A] to [Role B]. Write None only if no role or system transition occurs."
      }}
    ]
  }}
}}

═══════════════════════════════════════════
SOP STEP RULES — follow every one strictly
═══════════════════════════════════════════
- Every step MUST have all five fields: action, outcome_or_decision, happy_path, unhappy_path, handoffs.
- Never produce a step with only a title. Never summarise steps into fewer fields.
- action: always names the role performing it and the tool used if applicable.
- happy_path: always ends with "Continue to **Step N – Title**." using the exact step number and title.
- unhappy_path: always ends with "Return to **Step N – Title**." or "remains in **Step N – Title**." Never leave blank.
- Decision steps: title ends with "(Yes/No)". happy_path = Yes outcome. unhappy_path = No outcome.
- Extract IMPLIED steps from conditionals, role mentions, tool mentions, and casual narration.
- Steps in strict chronological order.

DECISION STEP RULES:
- Create a Decision step for EVERY conditional in the transcript.
- Never merge two conditionals into one Decision node.
- Every Decision step MUST have both a Yes branch (happy_path) and No branch (unhappy_path).

Transcript:
{transcript}
"""


def build_diagram_prompt(transcript: str) -> str:
    return f"""
You are a process documentation expert. Analyze the transcript and return ONLY this JSON — no extra text, no markdown fences.

{{
  "diagram": {{
    "lanes": [
      {{
        "id": "lane-1",
        "label": "Role name",
        "color": "#e8f4ff"
      }}
    ],
    "nodes": [
      {{
        "id": "node-1",
        "type": "process",
        "label": "Short step title",
        "description": "One-line detail shown on hover/expand",
        "laneId": "lane-1",
        "order": 1,
        "step_number": "01"
      }}
    ],
    "edges": [
      {{
        "id": "edge-1",
        "source": "node-1",
        "target": "node-2",
        "label": "",
        "type": "default"
      }}
    ]
  }}
}}

═══════════════════════════════════════════
DIAGRAM RULES — follow every one strictly
═══════════════════════════════════════════

LANES:
- One lane per distinct role/actor identified in the transcript.
- Use these color options cycling in order: "#e8f4ff", "#e8f9f0", "#fff4e8", "#f3e8ff", "#fef3c7", "#fce7f3".
- Lane IDs: "lane-1", "lane-2", "lane-3", etc.
- Lane label must match exactly the role name used in the SOP roles list.

NODE TYPES:
- "start"    — the trigger that begins the process (exactly one).
- "end"      — the terminal state (one or more allowed).
- "process"  — a standard action step.
- "decision" — a Yes/No branching point. Label must end with "?".

NODES:
- id: "node-1", "node-2", ... sequential across the whole diagram.
- laneId: must match one of the declared lane ids. Place the node in the lane of the role performing the action.
- order: horizontal position within the lane (1, 2, 3...). Used for layout.
- label: 2–6 words, verb-led (e.g. "Review invoice", "Approve request?").
- description: one sentence with detail such as the tool used or artifact produced.
- step_number: zero-padded string (e.g. "01", "02") matching the SOP step number when the node corresponds to an SOP step. Use "" for start/end nodes.

EDGES:
- Every node (except end nodes) must have at least one outgoing edge.
- Decision nodes must have EXACTLY two outgoing edges, labeled "Yes" and "No" respectively.
- Loop-backs are allowed — set target to an earlier node id and use type "loopback".
- type: "default" for normal flow, "loopback" for backward edges.
- Edge ids: "edge-1", "edge-2", ... sequential.

ORDER RULES:
- The start node always has order=1 in its lane.
- Nodes within a lane increment order based on flow sequence.
- Handoffs between lanes are represented purely as edges (no handoff node needed).

Transcript:
{transcript}
"""


def build_bottlenecks_prompt(transcript: str) -> str:
    return f"""
You are a process documentation expert. Analyze the transcript and return ONLY this JSON — no extra text, no markdown fences.

{{
  "bottlenecks": [
    {{
      "type": "Signal type name",
      "step_number": "03",
      "step_title": "Title of the affected step",
      "metrics": {{
        "transcript_mentions": 0,
        "rework_rate": "N/A",
        "times_delayed": 0
      }},
      "description": "2–3 sentences explaining exactly why this is a bottleneck.",
      "suggestions": [
        "Concrete actionable suggestion 1",
        "Concrete actionable suggestion 2",
        "Concrete actionable suggestion 3"
      ]
    }}
  ]
}}

═══════════════════════════════════════════
BOTTLENECK DETECTION RULES — MANDATORY
═══════════════════════════════════════════
You MUST scan every step against ALL 9 signal types below.

SIGNAL TYPES:
1. REPETITION — step mentioned more than twice, rework implied.
2. WAIT / HANDOFF DELAY — "waiting on", "haven't heard back", "still pending", "chasing".
3. MANUAL EFFORT — "we do it manually", "copy paste", "spreadsheet", "takes a while".
4. CONFUSION / AMBIGUITY — "not sure who owns this", "it varies", "I thought X did it".
5. VOLUME SPIKES — "backlog", "can't keep up", "overwhelmed", "piles up".
6. TOOL FRICTION — "the system doesn't", "we have to work around", "it breaks when".
7. ESCALATION PATTERNS — steps regularly escalated, re-approved, or sent back.
8. ABSENCE OF VALIDATION ← STRUCTURAL — flag ANY handoff with no confirmation gate or checklist.
9. ABSENCE OF SLA / TIMELINE ← STRUCTURAL — flag ANY waiting step with no turnaround time defined.

FOR EACH BOTTLENECK:
- type: exact signal name from the list above.
- step_number: zero-padded string (e.g. "03").
- transcript_mentions: count of direct and indirect references.
- rework_rate: % estimate if rework implied, else "N/A".
- times_delayed: explicit count. Absence signals: write 0.
- description: 2–3 sentences with transcript evidence or structural analysis.
- suggestions: exactly 3 concrete suggestions naming specific artifacts.

If no signals found: "bottlenecks": []

Transcript:
{transcript}
"""


# ─────────────────────────────────────────
#  SOP DOCUMENT PROMPTS  (new)
# ─────────────────────────────────────────

def build_restructure_sop_prompt(sop_text: str) -> str:
    """
    Maps an existing SOP document into our JSON schema.

    FAITHFULNESS RULES (baked in):
    - Copy step titles and action text verbatim — no paraphrasing.
    - If a field has no equivalent in the source, leave it empty string or [].
    - Never invent roles, steps, decisions, or SLAs not present in the source.
    - Step count must equal the original — no merging or splitting.
    - Any text that doesn't map to a schema field goes into the nearest step's action field.
    """
    return f"""
You are a process documentation specialist. Your task is to map an existing SOP document into a structured JSON schema.

CRITICAL FAITHFULNESS RULES — violating any of these is an error:
1. COPY exact wording from the source. Do NOT paraphrase, summarise, or rewrite any field.
2. PRESERVE every step — do not merge, split, reorder, or drop any step.
3. If the source does not contain a value for a field (e.g. no happy_path is defined), leave it as "" or [] — never invent content.
4. If a role is not explicitly named in the source, use "Unspecified" — never assume or invent a role.
5. Any source text that does not clearly map to a schema field must go into the nearest step's "action" field — never discard text.
6. The step "number" must match the original numbering in the source document exactly.

Return ONLY this JSON — no extra text, no markdown fences:

{{
  "sop": {{
    "title": "Exact title from the document",
    "generated_from": "sop_document",
    "date": "Date found in the document, or empty string if not present",
    "version": "Version found in the document, or empty string if not present",
    "participants": 0,
    "tags": [],
    "purpose": "Exact purpose/objective text from the document. Empty string if not present.",
    "scope": "Exact scope text from the document. Empty string if not present.",
    "roles": [
      {{
        "name": "Exact role name as written in the document",
        "responsibility": "Exact responsibility text as written in the document"
      }}
    ],
    "prerequisites": [
      "Exact prerequisite text as written"
    ],
    "start_state": "Exact trigger/start condition text from the document. Empty string if not present.",
    "end_state": "Exact completion condition text from the document. Empty string if not present.",
    "steps": [
      {{
        "number": 1,
        "title": "Exact step title copied verbatim from the document",
        "action": "Exact action text copied verbatim. Include all sub-bullets joined with semicolons.",
        "outcome_or_decision": "Exact outcome or decision text from the document. Empty string if not stated.",
        "happy_path": "Exact success path text from the document. Empty string if not stated.",
        "unhappy_path": "Exact failure/exception text from the document. Empty string if not stated.",
        "handoffs": "Exact handoff text from the document. Write None if no handoff is mentioned."
      }}
    ]
  }}
}}

SOP DOCUMENT TO RESTRUCTURE:
{sop_text}
"""


def build_diagram_from_sop_prompt(sop_text: str) -> str:
    """
    Derives a swimlane diagram from a structured SOP document.
    Roles → lanes. Steps → nodes. Flow text → edges.
    Only uses what is explicitly written in the source.
    """
    return f"""
You are a process documentation expert. Read the SOP document below and return ONLY a diagram JSON — no extra text, no markdown fences.

FAITHFULNESS RULES:
1. Lanes must come from roles explicitly named in the SOP. Do not invent roles.
2. Node labels must be short (2–6 words) verb-led versions of the exact step titles — no new wording.
3. Edges must follow the sequence and decision branches as written. Do not infer or add connections not stated.
4. If a step's owner/role is not stated, assign it to lane "Unspecified".

{{
  "diagram": {{
    "lanes": [
      {{
        "id": "lane-1",
        "label": "Exact role name from SOP",
        "color": "#e8f4ff"
      }}
    ],
    "nodes": [
      {{
        "id": "node-1",
        "type": "process",
        "label": "Short verb-led label (2–6 words)",
        "description": "Exact one-line excerpt from the SOP step action",
        "laneId": "lane-1",
        "order": 1,
        "step_number": "01"
      }}
    ],
    "edges": [
      {{
        "id": "edge-1",
        "source": "node-1",
        "target": "node-2",
        "label": "",
        "type": "default"
      }}
    ]
  }}
}}

DIAGRAM RULES:
- One lane per distinct role/actor named in the SOP.
- Lane colors cycling: "#e8f4ff", "#e8f9f0", "#fff4e8", "#f3e8ff", "#fef3c7", "#fce7f3".
- Node types: "start" (one only), "end" (one or more), "process" (standard step), "decision" (Yes/No branch, label ends with "?").
- Decision nodes must have EXACTLY two outgoing edges labeled "Yes" and "No".
- step_number: zero-padded string matching the SOP step number (e.g. "01"). Use "" for start/end nodes.
- Loop-back edges: type "loopback".
- Every non-end node must have at least one outgoing edge.

SOP DOCUMENT:
{sop_text}
"""


def build_bottlenecks_from_sop_prompt(sop_text: str) -> str:
    """
    Detects bottleneck signals in a written SOP document.

    Key difference from transcript bottlenecks:
    - No "mentions" count (it's a document, not a conversation).
    - Signals are STRUCTURAL: missing SLAs, missing owners, ambiguous steps,
      absent validation gates, manual-effort language in the procedure text.
    - transcript_mentions is set to 0; rework_rate and times_delayed reflect
      language in the document (e.g. "retry", "re-submit", "escalate").
    """
    return f"""
You are a process improvement expert. Read the SOP document below and identify bottlenecks.

IMPORTANT — you are reading a WRITTEN DOCUMENT, not a conversation transcript.
- Do not invent metrics that aren't in the document.
- Only flag what is explicitly stated or structurally absent in the text.
- Set transcript_mentions to 0 for all bottlenecks (there is no transcript).
- For rework_rate: use "N/A" unless the document explicitly mentions rework, retry, or resubmission.
- For times_delayed: use 0 unless the document mentions specific delay counts.

Return ONLY this JSON — no extra text, no markdown fences:

{{
  "bottlenecks": [
    {{
      "type": "Signal type name",
      "step_number": "03",
      "step_title": "Exact step title from the SOP",
      "metrics": {{
        "transcript_mentions": 0,
        "rework_rate": "N/A",
        "times_delayed": 0
      }},
      "description": "2–3 sentences quoting or closely referencing the SOP text that reveals the issue.",
      "suggestions": [
        "Concrete actionable suggestion 1",
        "Concrete actionable suggestion 2",
        "Concrete actionable suggestion 3"
      ]
    }}
  ]
}}

SIGNAL TYPES TO SCAN FOR:
1. MANUAL EFFORT — words like "manually", "copy", "paste", "email", "spreadsheet", "print", "fax".
2. AMBIGUOUS OWNERSHIP — step has no named role, or uses vague language like "someone", "the team", "as appropriate".
3. ABSENCE OF SLA / TIMELINE — a waiting or approval step with no turnaround time defined.
4. ABSENCE OF VALIDATION — a handoff step with no checklist, sign-off, or confirmation gate described.
5. ESCALATION PATTERN — the document describes steps that loop back, require re-approval, or escalate frequently.
6. TOOL FRICTION — language like "the system may not", "workaround", "manually check", "export and re-import".
7. MISSING EXCEPTION HANDLING — a step describes the happy path only with no stated action for failures or edge cases.
8. REDUNDANT STEPS — two consecutive steps that appear to perform the same action or produce the same output.
9. DEPENDENCY WITHOUT OWNER — the step depends on an external party or system with no contact, SLA, or fallback defined.

For each bottleneck found:
- type: exact signal name from the list above.
- step_number: zero-padded string of the affected step (e.g. "03"). Use "00" if the issue is in the header/preamble.
- description: quote or closely reference the exact SOP text that reveals the issue. Do not paraphrase beyond what is necessary.
- suggestions: exactly 3 concrete, actionable suggestions.

If no signals found: "bottlenecks": []

SOP DOCUMENT:
{sop_text}
"""


# ─────────────────────────────────────────
#  ENDPOINTS
# ─────────────────────────────────────────

@app.post("/analyse")
async def analyse_transcript(
    file: UploadFile = File(None),
    text: str = Form(None)
):
    """
    Transcript endpoint — handles Zoom .vtt and .txt files only.
    Runs all three transcript prompts (SOP, diagram, bottlenecks) in parallel.
    For .docx / .pdf SOP documents use /analyse-sop instead.
    """
    if file and file.filename:
        fname = (file.filename or "").lower()
        if fname.endswith(".docx") or fname.endswith(".pdf"):
            return {
                "error": (
                    "This looks like a document file. "
                    "Please use the SOP document upload path for .docx and .pdf files."
                )
            }
        transcript = await extract_text_from_file(file)
    elif text:
        transcript = text
    else:
        return {"error": "Please provide a file or paste transcript text."}

    try:
        sop_result, diagram_result, bottleneck_result = await asyncio.gather(
            call_openai(build_sop_prompt(transcript)),
            call_openai(build_diagram_prompt(transcript)),
            call_openai(build_bottlenecks_prompt(transcript)),
        )

        sop_data        = clean_json(sop_result)
        diagram_data    = clean_json(diagram_result)
        bottleneck_data = clean_json(bottleneck_result)

        bottleneck_step_numbers = {
            b["step_number"] for b in bottleneck_data.get("bottlenecks", [])
        }
        for node in diagram_data.get("diagram", {}).get("nodes", []):
            node["flag"] = node.get("step_number", "") in bottleneck_step_numbers

        return {
            **sop_data,
            **diagram_data,
            **bottleneck_data,
        }

    except json.JSONDecodeError:
        return {"error": "Failed to parse AI response. Please try again."}
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"error": str(e)}


@app.post("/analyse-sop")
async def analyse_sop_document(
    file: UploadFile = File(None),
    text: str = Form(None),
    intents: str = Form("restructure,diagram,bottlenecks")
):
    """
    New endpoint for uploaded SOP documents (.docx, .pdf, .txt).

    intents: comma-separated list of what the user wants.
             Values: "restructure", "diagram", "bottlenecks"
             Default: all three.

    Returns only the outputs that were requested.
    """
    if file and file.filename:
        try:
            sop_text = await extract_text_from_file(file)
        except RuntimeError as e:
            return {"error": str(e)}
    elif text:
        sop_text = text
    else:
        return {"error": "Please provide a file or paste SOP text."}

    if not sop_text.strip():
        return {"error": "Could not extract any text from the uploaded file."}

    requested = {i.strip() for i in intents.split(",")}

    # Build only the tasks that were requested
    tasks = {}
    if "restructure" in requested:
        tasks["sop"] = call_openai(build_restructure_sop_prompt(sop_text))
    if "diagram" in requested:
        tasks["diagram"] = call_openai(build_diagram_from_sop_prompt(sop_text))
    if "bottlenecks" in requested:
        tasks["bottlenecks"] = call_openai(build_bottlenecks_from_sop_prompt(sop_text))

    if not tasks:
        return {"error": "No valid intents specified. Use: restructure, diagram, bottlenecks"}

    try:
        results = await asyncio.gather(*tasks.values())
        parsed  = {key: clean_json(res) for key, res in zip(tasks.keys(), results)}

        response = {}

        if "sop" in parsed:
            response.update(parsed["sop"])                    # adds { "sop": { ... } }

        if "diagram" in parsed:
            response.update(parsed["diagram"])                # adds { "diagram": { ... } }

        if "bottlenecks" in parsed:
            bottleneck_data = parsed["bottlenecks"]
            response.update(bottleneck_data)                  # adds { "bottlenecks": [...] }

            # Flag bottleneck nodes on the diagram (if both were requested)
            if "diagram" in response:
                flagged = {
                    b["step_number"]
                    for b in bottleneck_data.get("bottlenecks", [])
                }
                for node in response.get("diagram", {}).get("nodes", []):
                    node["flag"] = node.get("step_number", "") in flagged

        return response

    except json.JSONDecodeError:
        return {"error": "Failed to parse AI response. Please try again."}
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"error": str(e)}


@app.get("/")
def root():
    return {"status": "ProcessLens API is running"}
